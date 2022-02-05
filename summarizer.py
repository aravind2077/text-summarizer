import docx
import pickle
def ReadingDoc(filename):
    transc = docx.Document(filename)
    fullText = []
    for i in transc.paragraphs:
        fullText.append(i.text)
    return '\n'.join(fullText)

def get_summary(chunks):
    filename = 'finalized_model.sav'
    summarizer = pickle.load(open(filename, 'rb'))
    res = summarizer(chunks, max_length=120, min_length=30, do_sample=False)
    summarized_text = ' '.join([summ['summary_text'] for summ in res])
        
    return summarized_text

print("Enter text in Transcript.docx")
body= ReadingDoc('Transcript.docx')
max_chunk = 500
body = body.replace('.', '.<eos>')
body = body.replace('?', '?<eos>')
body = body.replace('!', '!<eos>')

sentences = body.split('<eos>')
current_chunk = 0 
chunks = []
for sentence in sentences:
    if len(chunks) == current_chunk + 1: 
        if len(chunks[current_chunk]) + len(sentence.split(' ')) <= max_chunk:
            chunks[current_chunk].extend(sentence.split(' '))
        else:
            current_chunk += 1
            chunks.append(sentence.split(' '))
    else:
        print(current_chunk)
        chunks.append(sentence.split(' '))

for chunk_id in range(len(chunks)):
    chunks[chunk_id] = ' '.join(chunks[chunk_id])

summary= get_summary(chunks)

doc= docx.Document()
doc.add_paragraph(summary)
doc.save('Summary.docx')
print('Summarised text is saved in Summary.docx')